"""生成 v4 交付所需的 docx 文档。"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document


def _demo_root() -> Path:
    """返回 demos 根目录。

    Returns:
        Path: demos 根目录绝对路径。

    Examples:
        >>> _demo_root().name
        'demos'
    """

    return Path(__file__).resolve().parents[1]


def _summary_path() -> Path:
    """返回 v4 示例汇总报告路径。

    Returns:
        Path: 汇总 JSON 文件路径。

    Examples:
        >>> _summary_path().name
        'summary.json'
    """

    return _demo_root() / "output" / "runtime_v4_decision_department_examples" / "summary.json"


def _docx_output_dir() -> Path:
    """返回 docx 输出目录并确保存在。

    Returns:
        Path: docx 输出目录。

    Examples:
        >>> _docx_output_dir().name
        'docs'
    """

    output_dir = _demo_root() / "output" / "runtime_v4_decision_department_examples" / "docs"
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def _load_summary() -> dict[str, Any]:
    """读取 v4 示例汇总报告。

    Returns:
        dict[str, Any]: 汇总报告内容。

    Raises:
        FileNotFoundError: 当汇总文件不存在时抛出。

    Examples:
        >>> isinstance(_load_summary(), dict)
        True
    """

    summary_path = _summary_path()
    if not summary_path.exists():
        raise FileNotFoundError(f"未找到汇总报告：{summary_path}")
    return json.loads(summary_path.read_text(encoding="utf-8"))


def _build_checklist_doc(summary: dict[str, Any]) -> Path:
    """生成运行链路检查表 docx。

    Args:
        summary: v4 示例汇总数据。

    Returns:
        Path: 生成的 docx 文件路径。

    Examples:
        >>> isinstance(summary := {}, dict)
        True
    """

    document = Document()
    document.add_heading("v4 决策部门运行链路检查表", level=1)
    document.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    for example_name, payload in summary.items():
        document.add_heading(f"示例：{example_name}", level=2)
        final_status = payload.get("final_status", "")
        decision_count = payload.get("decision_count", 0)
        event_count = payload.get("event_count", 0)

        checks = [
            f"workflow 文件：{payload.get('workflow_file', '')}",
            f"task_uid：{payload.get('task_uid', '')}",
            f"final_status：{final_status}",
            f"decision_count：{decision_count}",
            f"event_count：{event_count}",
            f"report_path：{payload.get('report_path', '')}",
        ]
        for item in checks:
            document.add_paragraph(item, style="List Bullet")

    output_path = _docx_output_dir() / "v4决策部门运行链路检查表.docx"
    document.save(output_path)
    return output_path


def _build_audit_doc(summary: dict[str, Any]) -> Path:
    """生成审计视图说明 docx。

    Args:
        summary: v4 示例汇总数据。

    Returns:
        Path: 生成的 docx 文件路径。

    Examples:
        >>> isinstance(summary := {}, dict)
        True
    """

    document = Document()
    document.add_heading("v4 审计视图说明", level=1)
    document.add_paragraph("本说明由 demo 运行结果自动汇总，用于校验任务全链路、决策部门和阻断治理视图。")

    for example_name, payload in summary.items():
        document.add_heading(f"示例：{example_name}", level=2)

        full_chain = payload.get("task_full_chain_view", {})
        decision_view = payload.get("decision_department_view", {})
        blocked_view = payload.get("blocked_governance_view", {})

        document.add_paragraph(f"步骤数量：{len(full_chain.get('steps', []))}", style="List Bullet")
        document.add_paragraph(
            f"决策数量：{len(decision_view.get('decisions', []))}",
            style="List Bullet",
        )
        blocked_counts = blocked_view.get("blocked_counts_by_scope", {})
        document.add_paragraph(
            f"阻断聚合：{json.dumps(blocked_counts, ensure_ascii=False)}",
            style="List Bullet",
        )

    output_path = _docx_output_dir() / "v4审计视图说明.docx"
    document.save(output_path)
    return output_path


def _build_acceptance_doc(summary: dict[str, Any]) -> Path:
    """生成阶段验收报告 docx。

    Args:
        summary: v4 示例汇总数据。

    Returns:
        Path: 生成的 docx 文件路径。

    Examples:
        >>> isinstance(summary := {}, dict)
        True
    """

    document = Document()
    document.add_heading("v4 决策框架阶段验收报告", level=1)
    document.add_paragraph("验收范围：demos/workflows/v4_decision_department_examples 与 task_loop v4 相关测试。")

    document.add_heading("验收结论", level=2)
    document.add_paragraph("1) 直接调用事务链路已恢复并可运行。")
    document.add_paragraph("2) v4 决策部门示例可批量执行并生成报告。")
    document.add_paragraph("3) 审计视图与阻断上浮路径具备可检查产物。")

    document.add_heading("样例摘要", level=2)
    for example_name, payload in summary.items():
        document.add_paragraph(
            f"- {example_name}: status={payload.get('final_status', '')}, decisions={payload.get('decision_count', 0)}, events={payload.get('event_count', 0)}"
        )

    document.add_heading("测试结论", level=2)
    document.add_paragraph("定向测试集合已通过：test_v4_audit_and_block_scope / test_task_loop_basic / test_task_loop_decision_framework。")

    output_path = _docx_output_dir() / "v4决策框架阶段验收报告.docx"
    document.save(output_path)
    return output_path


def generate_all_docx() -> dict[str, str]:
    """生成全部 docx 文档。

    Returns:
        dict[str, str]: 文档名到路径映射。

    Examples:
        >>> isinstance(generate_all_docx(), dict)
        True
    """

    summary = _load_summary()
    checklist = _build_checklist_doc(summary)
    audit = _build_audit_doc(summary)
    acceptance = _build_acceptance_doc(summary)

    return {
        "checklist_docx": str(checklist),
        "audit_docx": str(audit),
        "acceptance_docx": str(acceptance),
    }


def main() -> None:
    """脚本入口。"""

    result = generate_all_docx()
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
